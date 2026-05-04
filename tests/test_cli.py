import json
import os
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document

from bruce_doc_converter import cli


REPO_ROOT = Path(__file__).resolve().parents[1]


class CliTests(unittest.TestCase):
    def run_cli(self, *args):
        return subprocess.run(
            [sys.executable, "-m", "bruce_doc_converter.cli", *args],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            encoding="utf-8",
            errors="replace",
        )

    def test_help_outputs_json(self):
        result = self.run_cli("--help-json")

        self.assertEqual(0, result.returncode, result.stderr)
        payload = json.loads(result.stdout)
        self.assertTrue(payload["success"])
        self.assertEqual("1.0", payload["schema_version"])
        self.assertIn("convert", payload["commands"])
        self.assertIn("batch", payload["commands"])

    def test_missing_command_outputs_json_failure(self):
        result = self.run_cli()

        self.assertEqual(1, result.returncode)
        payload = json.loads(result.stdout)
        self.assertFalse(payload["success"])
        self.assertEqual("USAGE_ERROR", payload["error_code"])

    def test_argparse_missing_positional_outputs_json(self):
        result = self.run_cli("convert")

        self.assertEqual(1, result.returncode)
        payload = json.loads(result.stdout)
        self.assertFalse(payload["success"])
        self.assertEqual("USAGE_ERROR", payload["error_code"])

    def test_argparse_invalid_choice_outputs_json(self):
        result = self.run_cli("convert", "file.docx", "--extract-images", "maybe")

        self.assertEqual(1, result.returncode)
        payload = json.loads(result.stdout)
        self.assertFalse(payload["success"])
        self.assertEqual("USAGE_ERROR", payload["error_code"])

    def test_dash_h_outputs_json(self):
        result = self.run_cli("-h")

        self.assertEqual(1, result.returncode)
        payload = json.loads(result.stdout)
        self.assertFalse(payload["success"])
        self.assertEqual("USAGE_ERROR", payload["error_code"])

    def test_argparse_error_emits_failure_exit_code(self):
        parser = cli._build_parser()

        with patch("bruce_doc_converter.cli._emit") as emit:
            with self.assertRaises(SystemExit) as raised:
                parser.parse_args(["convert"])

        self.assertEqual(1, raised.exception.code)
        self.assertEqual(1, emit.call_args.args[1])

    def test_normalize_result_uses_structured_error_code(self):
        payload = cli._normalize_single_result(
            "/tmp/input.docx",
            {"success": False, "error_code": "FILE_NOT_FOUND", "error": "upstream message changed"},
        )

        self.assertEqual("FILE_NOT_FOUND", payload["error_code"])

    def test_normalize_result_uses_unknown_for_missing_extension(self):
        payload = cli._normalize_single_result(
            "/tmp/no-extension",
            {"success": False, "error": "不支持的文件格式"},
        )

        self.assertEqual("unknown", payload["input_format"])


class CliConvertTests(unittest.TestCase):
    def run_cli(self, *args, cwd=None):
        env = os.environ.copy()
        env["PYTHONPATH"] = os.pathsep.join(
            [str(REPO_ROOT), env["PYTHONPATH"]] if env.get("PYTHONPATH") else [str(REPO_ROOT)]
        )
        return subprocess.run(
            [sys.executable, "-m", "bruce_doc_converter.cli", *args],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            encoding="utf-8",
            errors="replace",
            cwd=cwd,
            env=env,
        )

    def test_convert_docx_outputs_protocol_v1(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp_path = Path(tmp_dir)
            input_path = tmp_path / "sample.docx"
            doc = Document()
            doc.add_paragraph("正文")
            doc.save(input_path)

            result = self.run_cli("convert", str(input_path))

            self.assertEqual(0, result.returncode, result.stderr)
            payload = json.loads(result.stdout)
            self.assertEqual("1.0", payload["schema_version"])
            self.assertTrue(payload["success"])
            self.assertEqual(str(input_path.resolve()), payload["input_path"])
            self.assertEqual("docx", payload["input_format"])
            self.assertEqual("markdown", payload["output_format"])
            self.assertIn("正文", payload["markdown_content"])
            self.assertEqual([], payload["warnings"])
            self.assertIn("extracted_images", payload)

    def test_convert_missing_file_outputs_error_code(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            missing = Path(tmp_dir) / "missing.docx"

            result = self.run_cli("convert", str(missing))

            self.assertEqual(1, result.returncode)
            payload = json.loads(result.stdout)
            self.assertFalse(payload["success"])
            self.assertEqual("FILE_NOT_FOUND", payload["error_code"])
            self.assertEqual(str(missing.resolve()), payload["input_path"])

    def test_convert_output_path_is_absolute_with_relative_output_dir(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp_path = Path(tmp_dir)
            input_path = tmp_path / "sample.docx"
            doc = Document()
            doc.add_paragraph("正文")
            doc.save(input_path)

            result = self.run_cli("convert", str(input_path), "--output-dir", "relative_out", cwd=tmp_dir)

            self.assertEqual(0, result.returncode, result.stderr)
            payload = json.loads(result.stdout)
            self.assertTrue(payload["success"])
            self.assertTrue(Path(payload["output_path"]).is_absolute(), payload["output_path"])

    def test_batch_outputs_protocol_v1(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp_path = Path(tmp_dir)
            input_path = tmp_path / "sample.docx"
            doc = Document()
            doc.add_paragraph("批量正文")
            doc.save(input_path)

            result = self.run_cli("batch", str(tmp_path))

            self.assertEqual(0, result.returncode, result.stderr)
            payload = json.loads(result.stdout)
            self.assertEqual("1.0", payload["schema_version"])
            self.assertTrue(payload["success"])
            self.assertEqual(1, payload["total"])
            self.assertEqual(1, payload["succeeded"])
            self.assertEqual(0, payload["failed"])
            self.assertEqual(1, len(payload["results"]))
            self.assertNotIn("file", payload["results"][0])
            self.assertEqual(
                payload["results"][0]["input_path"],
                payload["results"][0]["result"]["input_path"],
            )

    def test_batch_with_output_dir_preserves_relative_directories_for_duplicate_names(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp_path = Path(tmp_dir)
            for subdir, text in (("a", "来自A"), ("b", "来自B")):
                input_dir = tmp_path / subdir
                input_dir.mkdir()
                doc = Document()
                doc.add_paragraph(text)
                doc.save(input_dir / "same.docx")

            output_dir = tmp_path / "out"
            result = self.run_cli("batch", str(tmp_path), "--output-dir", str(output_dir))

            self.assertEqual(0, result.returncode, result.stderr)
            payload = json.loads(result.stdout)
            self.assertTrue(payload["success"])
            output_paths = {item["result"]["output_path"] for item in payload["results"]}
            self.assertEqual(2, len(output_paths))
            self.assertEqual("来自A", (output_dir / "a" / "same.md").read_text(encoding="utf-8"))
            self.assertEqual("来自B", (output_dir / "b" / "same.md").read_text(encoding="utf-8"))
